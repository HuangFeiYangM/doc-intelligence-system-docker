"""
Tests for DocumentParser module.
"""
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook

from app.services.document_parser import DocumentParser, DocumentParserError
from app.models import DocumentType


class TestDocumentParser:
    """Test cases for DocumentParser."""

    def test_get_document_type_pdf(self):
        """Test detecting PDF document type."""
        assert DocumentParser.get_document_type("test.pdf") == DocumentType.PDF
        assert DocumentParser.get_document_type("/path/to/file.PDF") == DocumentType.PDF

    def test_get_document_type_word(self):
        """Test detecting Word document type."""
        assert DocumentParser.get_document_type("test.docx") == DocumentType.WORD
        assert DocumentParser.get_document_type("test.doc") == DocumentType.WORD
        assert DocumentParser.get_document_type("/path/to/file.DOCX") == DocumentType.WORD

    def test_get_document_type_excel(self):
        """Test detecting Excel document type."""
        assert DocumentParser.get_document_type("test.xlsx") == DocumentType.EXCEL
        assert DocumentParser.get_document_type("test.xls") == DocumentType.EXCEL

    def test_get_document_type_unsupported(self):
        """Test detecting unsupported document type."""
        with pytest.raises(DocumentParserError) as exc_info:
            DocumentParser.get_document_type("test.txt")
        assert "Unsupported file extension" in str(exc_info.value)

    def test_parse_word(self, temp_dir):
        """Test parsing Word document."""
        # Create a test Word document
        doc_path = temp_dir / "test.docx"
        doc = DocxDocument()
        doc.add_paragraph("Test Paragraph 1")
        doc.add_paragraph("Test Paragraph 2")
        doc.add_paragraph("")
        doc.add_paragraph("Test Paragraph 3")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"

        doc.save(doc_path)

        # Parse the document
        text = DocumentParser.parse_word(doc_path)

        assert "Test Paragraph 1" in text
        assert "Test Paragraph 2" in text
        assert "Test Paragraph 3" in text
        assert "Cell 1" in text
        assert "Cell 2" in text

    def test_parse_word_not_found(self):
        """Test parsing non-existent Word document."""
        with pytest.raises(DocumentParserError) as exc_info:
            DocumentParser.parse_word("/nonexistent/document.docx")
        assert "Failed to parse" in str(exc_info.value)

    def test_parse_excel(self, temp_dir):
        """Test parsing Excel document."""
        # Create a test Excel file
        excel_path = temp_dir / "test.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"] = "Header1"
        ws["B1"] = "Header2"
        ws["A2"] = "Data1"
        ws["B2"] = "Data2"

        # Add second sheet
        ws2 = wb.create_sheet("Sheet2")
        ws2["A1"] = "Sheet2 Header"

        wb.save(excel_path)

        # Parse the document
        text = DocumentParser.parse_excel(excel_path)

        assert "Sheet: Sheet1" in text
        assert "Sheet: Sheet2" in text
        assert "Header1" in text
        assert "Data1" in text

    def test_parse_excel_not_found(self):
        """Test parsing non-existent Excel document."""
        with pytest.raises(DocumentParserError) as exc_info:
            DocumentParser.parse_excel("/nonexistent/document.xlsx")
        assert "Failed to parse" in str(exc_info.value)

    def test_extract_text_file_not_found(self):
        """Test extracting text from non-existent file."""
        with pytest.raises(DocumentParserError) as exc_info:
            DocumentParser.extract_text("/nonexistent/file.pdf")
        assert "File not found" in str(exc_info.value)

    def test_extract_text_with_metadata_word(self, temp_dir):
        """Test extracting text with metadata from Word document."""
        # Create a test Word document
        doc_path = temp_dir / "test.docx"
        doc = DocxDocument()
        doc.add_paragraph("Test content")
        doc.save(doc_path)

        # Extract with metadata
        result = DocumentParser.extract_text_with_metadata(doc_path)

        assert "text" in result
        assert "file_name" in result
        assert "file_size" in result
        assert "document_type" in result
        assert result["file_name"] == "test.docx"
        assert result["document_type"] == "word"
        assert result["file_size"] > 0

    def test_extract_text_with_metadata_excel(self, temp_dir):
        """Test extracting text with metadata from Excel document."""
        # Create a test Excel file
        excel_path = temp_dir / "test.xlsx"
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "Test"
        wb.save(excel_path)

        # Extract with metadata
        result = DocumentParser.extract_text_with_metadata(excel_path)

        assert result["file_name"] == "test.xlsx"
        assert result["document_type"] == "excel"


class TestDocumentParserEdgeCases:
    """Edge case tests for DocumentParser."""

    def test_empty_word_document(self, temp_dir):
        """Test parsing empty Word document."""
        doc_path = temp_dir / "empty.docx"
        doc = DocxDocument()
        doc.save(doc_path)

        text = DocumentParser.parse_word(doc_path)
        assert text.strip() == ""

    def test_empty_excel_document(self, temp_dir):
        """Test parsing empty Excel document."""
        excel_path = temp_dir / "empty.xlsx"
        wb = Workbook()
        wb.save(excel_path)

        text = DocumentParser.parse_excel(excel_path)
        assert "Sheet" in text

    def test_word_with_only_tables(self, temp_dir):
        """Test parsing Word document with only tables."""
        doc_path = temp_dir / "tables_only.docx"
        doc = DocxDocument()

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Value 1"
        table.cell(1, 1).text = "Value 2"

        doc.save(doc_path)

        text = DocumentParser.parse_word(doc_path)
        assert "Header 1" in text
        assert "Value 1" in text