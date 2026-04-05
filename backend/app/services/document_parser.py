"""
Document parser module for extracting text from various document formats.
"""
import logging
import os
from pathlib import Path
from typing import Optional, Union

import pandas as pd
import pdfplumber
from docx import Document as DocxDocument

from app.config import get_settings
from app.models import DocumentType

# Configure logger
logger = logging.getLogger(__name__)

settings = get_settings()


class DocumentParserError(Exception):
    """Exception raised for document parsing errors."""
    pass


class DocumentParser:
    """Parser for extracting text from PDF, Word, and Excel documents."""

    @staticmethod
    def get_document_type(file_path: Union[str, Path]) -> DocumentType:
        """Determine document type from file extension."""
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return DocumentType.PDF
        elif ext in [".docx", ".doc"]:
            return DocumentType.WORD
        elif ext in [".xlsx", ".xls"]:
            return DocumentType.EXCEL
        else:
            raise DocumentParserError(f"Unsupported file extension: {ext}")

    @staticmethod
    def parse_pdf(file_path: Union[str, Path]) -> str:
        """Extract text from PDF file.

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text content

        Raises:
            DocumentParserError: If parsing fails
        """
        file_path = Path(file_path)
        logger.info(f"Starting PDF parsing: {file_path}")

        try:
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                logger.debug(f"PDF has {page_count} pages")

                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")

            result = "\n".join(text_parts)
            logger.info(f"PDF parsing completed: {file_path}, extracted {len(result)} characters")
            return result

        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}", exc_info=True)
            raise DocumentParserError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def parse_word(file_path: Union[str, Path]) -> str:
        """Extract text from Word document.

        Args:
            file_path: Path to the Word document

        Returns:
            Extracted text content

        Raises:
            DocumentParserError: If parsing fails
        """
        file_path = Path(file_path)
        logger.info(f"Starting Word document parsing: {file_path}")

        try:
            doc = DocxDocument(file_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            result = "\n".join(text_parts)
            logger.info(f"Word document parsing completed: {file_path}, extracted {len(result)} characters")
            return result

        except Exception as e:
            logger.error(f"Failed to parse Word document {file_path}: {e}", exc_info=True)
            raise DocumentParserError(f"Failed to parse Word document: {str(e)}")

    @staticmethod
    def parse_excel(file_path: Union[str, Path]) -> str:
        """Extract text from Excel file.

        Args:
            file_path: Path to the Excel file

        Returns:
            Extracted text content (CSV-like format)

        Raises:
            DocumentParserError: If parsing fails
        """
        file_path = Path(file_path)
        logger.info(f"Starting Excel parsing: {file_path}")

        try:
            # Read all sheets using context manager to ensure file is closed
            text_parts = []

            with pd.ExcelFile(file_path) as excel_file:
                sheet_names = excel_file.sheet_names
                logger.debug(f"Excel has {len(sheet_names)} sheets: {sheet_names}")

                for sheet_name in sheet_names:
                    try:
                        df = pd.read_excel(excel_file, sheet_name=sheet_name)
                        text_parts.append(f"Sheet: {sheet_name}")
                        text_parts.append(df.to_string(index=False))
                        text_parts.append("-" * 50)
                    except Exception as e:
                        logger.warning(f"Failed to read sheet '{sheet_name}': {e}")

            result = "\n".join(text_parts)
            logger.info(f"Excel parsing completed: {file_path}, extracted {len(result)} characters")
            return result

        except Exception as e:
            logger.error(f"Failed to parse Excel file {file_path}: {e}", exc_info=True)
            raise DocumentParserError(f"Failed to parse Excel file: {str(e)}")

    @classmethod
    def extract_text(cls, file_path: Union[str, Path]) -> str:
        """Unified entry point for extracting text from any supported document.

        Args:
            file_path: Path to the document file

        Returns:
            Extracted text content

        Raises:
            DocumentParserError: If file not found or parsing fails
        """
        file_path = Path(file_path)
        logger.info(f"Starting text extraction: {file_path}")

        # Path fallback mechanism: if file doesn't exist at the given path,
        # try alternative paths (handles path mismatches between API and Celery worker)
        if not file_path.exists():
            original_path = file_path
            # Try to find the file in alternative locations
            alternative_paths = [
                # Try with settings.UPLOAD_DIR
                settings.UPLOAD_DIR / file_path.name,
                # Try prepending /app (common Docker path issue)
                Path("/app") / file_path.relative_to(file_path.anchor) if file_path.is_absolute() else Path("/app") / file_path,
            ]

            for alt_path in alternative_paths:
                if alt_path.exists():
                    logger.warning(f"File not found at {original_path}, using alternative path: {alt_path}")
                    file_path = alt_path
                    break

        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            raise DocumentParserError(f"File not found: {file_path}")

        if not os.access(file_path, os.R_OK):
            logger.error(f"File not readable: {file_path}")
            raise DocumentParserError(f"File not readable: {file_path}")

        try:
            doc_type = cls.get_document_type(file_path)
            logger.debug(f"Document type detected: {doc_type}")

            if doc_type == DocumentType.PDF:
                result = cls.parse_pdf(file_path)
            elif doc_type == DocumentType.WORD:
                result = cls.parse_word(file_path)
            elif doc_type == DocumentType.EXCEL:
                result = cls.parse_excel(file_path)
            else:
                raise DocumentParserError(f"Unsupported document type: {doc_type}")

            logger.info(f"Text extraction completed: {file_path}, {len(result)} characters")
            return result

        except DocumentParserError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error extracting text from {file_path}: {e}", exc_info=True)
            raise DocumentParserError(f"Failed to extract text: {str(e)}")

    @classmethod
    def extract_text_with_metadata(cls, file_path: Union[str, Path]) -> dict:
        """Extract text with metadata.

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary containing text content and metadata
        """
        file_path = Path(file_path)
        logger.info(f"Extracting text with metadata: {file_path}")

        text = cls.extract_text(file_path)
        doc_type = cls.get_document_type(file_path)
        file_stat = file_path.stat()

        metadata = {
            "text": text,
            "file_name": file_path.name,
            "file_size": file_stat.st_size,
            "file_modified": file_stat.st_mtime,
            "document_type": doc_type.value,
        }

        logger.info(f"Metadata extraction completed: {file_path}")
        return metadata
